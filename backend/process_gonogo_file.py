import json
import logging
import os
import time
import asyncio
import tempfile
import uuid
from datetime import datetime

from botocore.exceptions import NoCredentialsError
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv
from openai import OpenAI, AsyncOpenAI
from backend.s3_config import put_object

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

load_dotenv()

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


def prepare_fine_tuning_data(results):
    fine_tuning_data = []
    for item in results:
        fine_tuning_data.append({
            "messages": [
                {"role": "system", "content": "Vous êtes un expert en bureau d'études, capable d'analyser des documents techniques et de fournir des analyses détaillées. BU: '', Métier / Société: '', Donneur d'ordres: '', Opportunité: '', Calendrier: { Date limite de remise des offres: '', Début de la prestation: '', Délai de validité des offres: '', Autres dates importantes: [] }, Critères d'attribution: [], Description de l'offre: { Durée: '', Synthèse Lot: '', CA TOTAL offensif: '', Missions générales: [], Matériels à disposition: [] }, Objet du marché: '', Périmètre de la consultation: '', Description des prestations: [], Exigences: [], Missions et compétences attendues: [], Profil des hôtes ou hôtesses d'accueil: { Qualités: [], Compétences nécessaires: [] }, Plages horaires: [], PSE: '', Formations: [], Intérêt pour le groupe: { Forces: [], Faiblesses: [], Opportunités: [], Menaces: [] }, Formule de révision des prix: ''"},
                {"role": "user", "content": f"Analysez le contenu suivant : {item}"},
                {"role": "assistant", "content": "Voici l'analyse du contenu fourni : [Insérez ici une analyse détaillée]"}
            ]
        })
    return fine_tuning_data

def save_jsonl(data, filename):
    with open(filename, 'w', encoding='utf-8') as f:
        for item in data:
            json.dump(item, f, ensure_ascii=False)
            f.write('\n')

def upload_file_for_fine_tuning(file_path):
    try:
        with open(file_path, 'rb') as file:
            response = client.files.create(file=file, purpose='fine-tune')
        return response.id
    except Exception as e:
        logger.error(f"Erreur lors du téléchargement du fichier : {str(e)}")
        return None

def start_fine_tuning(file_id):
    try:
        response = client.fine_tuning.jobs.create(training_file=file_id, model="gpt-3.5-turbo")
        return response.id
    except Exception as e:
        logger.error(f"Erreur lors du démarrage du fine-tuning : {str(e)}")
        return None

def split_content(content, max_tokens=4000):
    if not isinstance(content, str):
        content = json.dumps(content, indent=2)

    words = content.split()
    chunks = []
    current_chunk = []
    current_length = 0

    for word in words:
        if current_length + len(word) + 1 > max_tokens:
            chunks.append(' '.join(current_chunk))
            current_chunk = [word]
            current_length = len(word)
        else:
            current_chunk.append(word)
            current_length += len(word) + 1

    if current_chunk:
        chunks.append(' '.join(current_chunk))

    return chunks

async def get_chatgpt_response_async(client, chunk):
    try:
        response = await client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "Vous êtes un assistant d'analyse de documents. Veuillez structurer votre réponse selon le format suivant, en remplissant uniquement les sections pour lesquelles vous avez des informations pertinentes :"},
                {"role": "user", "content": f"""
                    Analysez le contenu suivant et fournissez les informations pertinentes pour chaque catégorie. Si vous n'avez pas d'information pour une catégorie, laissez-la vide.

                    {chunk}

                    BU :
                    Métier / Société :
                    Donneur d'ordres :
                    Opportunité :
                    Date limite de remise des offres :
                    Début de la prestation :
                    Délai de validité des offres :
                    Objet du marché :
                    Périmètre de la consultation :
                    PSE :
                    Formule de révision des prix :
                    Critères d'attribution :
                    Description des prestations :
                    Exigences :
                    Missions et compétences attendues :
                    Qualités des hôtes ou hôtesses :
                    Compétences nécessaires :
                    Plages horaires :
                    Formations :
                    Forces :
                    Faiblesses :
                    Opportunités :
                    Menaces :

                    Assurez-vous de ne remplir que les catégories pour lesquelles vous avez des informations spécifiques et pertinentes.
                    """
                }
            ]
        )
        chatgpt_response = response.choices[0].message.content
        logger.debug(f"Réponse brute de ChatGPT : {chatgpt_response}")
        return {"analyse": chatgpt_response}
    except Exception as e:
        logger.error(f"Erreur lors de l'appel à l'API ChatGPT : {str(e)}", exc_info=True)
        return {"erreur": str(e)}

async def process_chunks(chunks):
    client = AsyncOpenAI()
    tasks = [get_chatgpt_response_async(client, chunk) for chunk in chunks]
    results = await asyncio.gather(*tasks)
    logger.info(f"Nombre de chunks traités : {len(results)}")
    return results

def upload_to_s3(local_file, s3_file):
    try:
        with open(local_file, 'rb') as file:
            file_data = file.read()
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            if put_object('jobpilot', s3_file, file_data, len(file_data), content_type):
                return s3_file
            else:
                logger.error(f"Échec de l'upload du fichier {local_file} sur MinIO")
                return s3_file
    except Exception as e:
        logger.error(f"Erreur lors de l'upload vers MinIO : {str(e)}")
        return False

def download_from_s3(s3_file, local_file):
    try:
        s3_client.download_file("jobpilot", s3_file, local_file)
        logger.info(f"Fichier {s3_file} téléchargé depuis S3 avec succès")
        return True
    except NoCredentialsError:
        logger.error("Identifiants non valides pour accéder à S3")
        return False

def create_word_document(merged_analysis):
    doc = Document()
    doc.add_heading('Analyse du document', 0)

    for key, value in merged_analysis.items():
        if isinstance(value, dict):
            doc.add_heading(key, level=1)
            for sub_key, sub_value in value.items():
                doc.add_paragraph(f"{sub_key}: {sub_value}")
        elif isinstance(value, list):
            doc.add_heading(key, level=1)
            for item in value:
                doc.add_paragraph(f"- {item}")
        else:
            doc.add_paragraph(f"{key}: {value}")

    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    logger.info(f"Document Word créé avec succès : {temp_file.name}")
    return temp_file.name

def split_results(results, max_tokens):
    parts = []
    current_part = []
    current_tokens = 0

    for result in results:
        result_tokens = len(json.dumps(result)) // 8
        if current_tokens + result_tokens > max_tokens:
            parts.append(current_part)
            current_part = []
            current_tokens = 0
        current_part.append(result)
        current_tokens += result_tokens

    if current_part:
        parts.append(current_part)

    return parts

def merge_analyses(analyses):
    merged = {
        "BU": "",
        "Métier / Société": "",
        "Donneur d'ordres": "",
        "Opportunité": "",
        "Calendrier": {
            "Date limite de remise des offres": "",
            "Début de la prestation": "",
            "Délai de validité des offres": "",
            "Autres dates importantes": []
        },
        "Critères d'attribution": [],
        "Description de l'offre": {
            "Durée": "",
            "Synthèse Lot": "",
            "CA TOTAL offensif": "",
            "Missions générales": [],
            "Matériels à disposition": []
        },
        "Objet du marché": "",
        "Périmètre de la consultation": "",
        "Description des prestations": [],
        "Exigences": [],
        "Missions et compétences attendues": [],
        "Profil des hôtes ou hôtesses d'accueil": {
            "Qualités": [],
            "Compétences nécessaires": []
        },
        "Plages horaires": [],
        "PSE": "",
        "Formations": [],
        "Intérêt pour le groupe": {
            "Forces": [],
            "Faiblesses": [],
            "Opportunités": [],
            "Menaces": []
        },
        "Formule de révision des prix": ""
    }

    for analysis in analyses:
        if isinstance(analysis, dict) and "analyse" in analysis:
            lines = analysis["analyse"].split("\n")
            current_key = ""
            for line in lines:
                if ":" in line:
                    key, value = line.split(":", 1)
                    key = key.strip()
                    value = value.strip()
                    if key in merged:
                        if isinstance(merged[key], str):
                            merged[key] = value
                        elif isinstance(merged[key], list):
                            if value:
                                merged[key].append(value)
                    elif key in merged["Calendrier"]:
                        merged["Calendrier"][key] = value
                    elif key in merged["Description de l'offre"]:
                        if isinstance(merged["Description de l'offre"][key], str):
                            merged["Description de l'offre"][key] = value
                        elif isinstance(merged["Description de l'offre"][key], list):
                            if value:
                                merged["Description de l'offre"][key].append(value)
                    elif key in merged["Profil des hôtes ou hôtesses d'accueil"]:
                        if value:
                            merged["Profil des hôtes ou hôtesses d'accueil"][key].append(value)
                    elif key in merged["Intérêt pour le groupe"]:
                        if value:
                            merged["Intérêt pour le groupe"][key].append(value)
                    current_key = key
                elif current_key and line.strip():
                    if isinstance(merged.get(current_key), list):
                        merged[current_key].append(line.strip())
                    elif current_key in merged["Intérêt pour le groupe"]:
                        merged["Intérêt pour le groupe"][current_key].append(line.strip())

    logger.info(f"Résultat final de merge_analyses : {merged}")
    return merged

async def process_gonogo_file(results):
    try:
        logger.info("Début du traitement des résultats")

        all_analyses = []
        for result in results:
            content = json.dumps(result, indent=2)
            chunks = split_content(content)
            
            chunk_analyses = await process_chunks(chunks)
            
            valid_analyses = [analysis["analyse"] for analysis in chunk_analyses if "erreur" not in analysis]
            if valid_analyses:
                merged_analysis = merge_analyses(valid_analyses)
                all_analyses.append(merged_analysis)

        final_analysis = merge_analyses(all_analyses)

        word_document_path = create_word_document(final_analysis)
        logger.info(f"Document Word créé avec succès : {word_document_path}")

        # Générer un nom de fichier unique
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_id = str(uuid.uuid4())[:8]
        filename = f"analyse_{timestamp}_{unique_id}.docx"
        
        s3_file_path = f"analyses/{filename}"
        
        upload_success = upload_to_s3(word_document_path, s3_file_path)
        if upload_success:
            logger.info(f"Document Word uploadé vers MinIO : {s3_file_path}")
        else:
            logger.error(f"Échec de l'upload du document Word vers MinIO")

        if word_document_path:
            os.remove(word_document_path)
            logger.info(f"Fichier temporaire {word_document_path} supprimé")

        return {
            "message": "Traitement terminé avec succès",
            "chatgpt_analysis": final_analysis,
            "word_document": upload_success
        }
    except Exception as e:
        logger.error(f"Erreur lors du traitement : {str(e)}", exc_info=True)
        return {"error": str(e)}

def read_jsonl_file(file_path):
    logging.info(f"Tentative de lecture du fichier : {file_path}")
    try:
        local_file = 'temp_jsonl_file.jsonl'
        if download_from_s3(file_path, local_file):
            with open(local_file, 'r') as file:
                content = [json.loads(line) for line in file]
            os.remove(local_file)
            logging.info(f"Fichier {file_path} lu avec succès depuis S3.")
            return content
        else:
            logging.error(f"Impossible de télécharger le fichier {file_path} depuis S3.")
    except Exception as e:
        logging.error(f"Erreur inattendue lors de la lecture du fichier {file_path} : {str(e)}")
    return None
